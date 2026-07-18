import os
import uuid
import hashlib
from datetime import datetime
from typing import List, Optional
from pathlib import Path

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from app.config import settings


class DocumentProcessor:
    SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx", ".csv"}

    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.chunk_size,
            chunk_overlap=settings.chunk_overlap,
            separators=["\n\n", "\n", ".", " ", ""],
        )

    def read_file(self, file_path: str) -> Optional[str]:
        ext = Path(file_path).suffix.lower()
        if ext == ".txt":
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                return f.read()
        elif ext == ".md":
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                return f.read()
        elif ext == ".pdf":
            return self._read_pdf(file_path)
        elif ext == ".docx":
            return self._read_docx(file_path)
        elif ext == ".csv":
            return self._read_csv(file_path)
        return None

    def _read_pdf(self, file_path: str) -> str:
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    def _read_docx(self, file_path: str) -> str:
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    def _read_csv(self, file_path: str) -> str:
        import csv
        rows = []
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            reader = csv.reader(f)
            for row in reader:
                rows.append(", ".join(row))
        return "\n".join(rows)

    def process_file(self, file_path: str, original_filename: str) -> List[Document]:
        content = self.read_file(file_path)
        if content is None:
            raise ValueError(f"Unsupported file type: {file_path}")

        file_hash = hashlib.md5(content.encode()).hexdigest()
        chunks = self.text_splitter.split_text(content)

        doc_id = str(uuid.uuid4())
        documents = []
        for i, chunk in enumerate(chunks):
            metadata = {
                "source": original_filename,
                "doc_id": doc_id,
                "chunk_index": i,
                "file_hash": file_hash,
                "uploaded_at": datetime.now().isoformat(),
            }
            documents.append(Document(page_content=chunk, metadata=metadata))

        return documents


processor = DocumentProcessor()
