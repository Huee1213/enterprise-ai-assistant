import os
import uuid
import hashlib
from datetime import datetime
from typing import List, Optional
from pathlib import Path

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from app.config import settings


class DocumentProcessor:
    SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx", ".csv"}

    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.chunk_size,
            chunk_overlap=settings.chunk_overlap,
            separators=["\n\n", "\n", ".", " ", ""],
        )

    def read_file(self, file_path: str) -> Optional[str]:
        ext = Path(file_path).suffix.lower()
        try:
            if ext in {".txt", ".md"}:
                with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                    return f.read()
            elif ext == ".pdf":
                return self._read_pdf(file_path)
            elif ext == ".docx":
                return self._read_docx(file_path)
            elif ext == ".csv":
                return self._read_csv(file_path)
        except Exception as e:
            raise ValueError(f"读取文件失败 ({ext}): {str(e)[:100]}")
        return None

    def _read_pdf(self, file_path: str) -> str:
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        return "\n".join(pages)

    def _read_docx(self, file_path: str) -> str:
        """Extract all text from DOCX including paragraphs, tables, headers, footers."""
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        parts = []

        # Body paragraphs
        for p in doc.paragraphs:
            t = p.text.strip()
            if t:
                parts.append(t)

        # Tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        # Headers and footers
        for section in doc.sections:
            for header in [section.header, section.first_page_header]:
                if header:
                    for p in header.paragraphs:
                        if p.text.strip():
                            parts.append(p.text.strip())
            for footer in [section.footer, section.first_page_footer]:
                if footer:
                    for p in footer.paragraphs:
                        if p.text.strip():
                            parts.append(p.text.strip())

        return "\n".join(parts)

    def _read_csv(self, file_path: str) -> str:
        import csv
        rows = []
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            reader = csv.reader(f)
            for row in reader:
                rows.append(", ".join(row))
        return "\n".join(rows)

    def process_file(self, file_path: str, original_filename: str, doc_id: Optional[str] = None) -> List[Document]:
        content = self.read_file(file_path)
        if not content:
            raise ValueError(f"文件内容为空或无法解析: {original_filename}")

        file_hash = hashlib.md5(content.encode()).hexdigest()
        chunks = self.text_splitter.split_text(content)

        doc_id = doc_id or str(uuid.uuid4())
        documents = []
        for i, chunk in enumerate(chunks):
            metadata = {
                "source": original_filename,
                "doc_id": doc_id,
                "chunk_index": i,
                "file_hash": file_hash,
                "uploaded_at": datetime.now().isoformat(),
            }
            documents.append(Document(page_content=chunk, metadata=metadata))

        return documents


processor = DocumentProcessor()
